"""
fill_template.py
----------------
Reads word_placeholder_excel.xlsx (row 1 = placeholder names, rows 2+ = values),
fills a Word template (.docx) by replacing [placeholder] tags,
and writes one .docx (and optionally .pdf) per row to the output folder.

The [File name] column controls the output filename suffix.
Output: [template filename] - [File name value].docx / .pdf

Placeholder format in the Word template must be exactly [name], no spaces.
Empty cells are allowed and will be replaced with an empty string.
"""

import os
import sys

# ── CONFIG ────────────────────────────────────────────────────────────────────
SCRIPT_DIR    = os.path.dirname(os.path.abspath(__file__))
PARENT_DIR    = os.path.dirname(SCRIPT_DIR)
EXCEL_PATH    = os.path.join(SCRIPT_DIR, "word_placeholder_excel.xlsx")
OUTPUT_FOLDER = os.path.join(PARENT_DIR, "output")
# ─────────────────────────────────────────────────────────────────────────────


def check_dependencies():
    packages = {"docx": "python-docx", "openpyxl": "openpyxl"}
    missing = []
    for import_name, install_name in packages.items():
        try:
            __import__(import_name)
        except ImportError:
            missing.append(install_name)
    if missing:
        print(f"Missing packages: {', '.join(missing)}")
        print(f"Run:  pip install {' '.join(missing)}")
        sys.exit(1)


# ── EXCEL ─────────────────────────────────────────────────────────────────────

def update_excel_headers(placeholder_names, log=print):
    """
    Called from the launcher after the user confirms placeholders.
    Wipes the Excel and writes a fresh header row.
    placeholder_names: list of str, e.g. ['company', 'address', 'contact']
    """
    import openpyxl

    if not placeholder_names:
        log("ERROR: No placeholder names provided.")
        return False

    # Check file is not open (write-test)
    try:
        with open(EXCEL_PATH, "a"):
            pass
    except IOError:
        log(f"ERROR: {os.path.basename(EXCEL_PATH)} is open in Excel. Close it and try again.")
        return False

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["File name"] + placeholder_names)
    wb.save(EXCEL_PATH)
    log(f"Excel updated — columns: File name, {', '.join(placeholder_names)}")
    log(f"Fill in your data starting from row 2, then press Run.")
    return True


def read_excel(log=print):
    """
    Returns (headers, rows).
    headers: list of str (row 1).
    rows: list of dicts {header: value}.
    Exits with an error message on any structural problem.
    """
    import openpyxl

    if not os.path.isfile(EXCEL_PATH):
        log(f"ERROR: Excel not found: {EXCEL_PATH}")
        sys.exit(1)

    # Check file is not open
    try:
        with open(EXCEL_PATH, "a"):
            pass
    except IOError:
        log(f"ERROR: {os.path.basename(EXCEL_PATH)} is open in Excel. Close it and try again.")
        sys.exit(1)

    wb = openpyxl.load_workbook(EXCEL_PATH, read_only=True, data_only=True)
    ws = wb.active
    all_rows = list(ws.iter_rows(values_only=True))
    wb.close()

    if not all_rows:
        log("ERROR: Excel is empty.")
        sys.exit(1)

    # Headers from row 1, strip trailing empty columns
    headers = [str(h).strip() if h is not None else "" for h in all_rows[0]]
    while headers and not headers[-1]:
        headers.pop()

    if not headers:
        log("ERROR: Header row is empty. Run 'Update Excel' first.")
        sys.exit(1)

    # Data rows (skip fully blank rows)
    rows = []
    for row_idx, raw_row in enumerate(all_rows[1:], start=2):
        values = list(raw_row[: len(headers)])
        # Pad if row is shorter than header count
        while len(values) < len(headers):
            values.append(None)

        # Skip rows that are entirely empty
        if all(v is None or str(v).strip() == "" for v in values):
            continue

        # Check for any empty cell
        row_dict = {}
        for col_idx, (h, v) in enumerate(zip(headers, values), start=1):
            cell_val = str(v).strip() if v is not None else ""
            row_dict[h] = cell_val

        rows.append(row_dict)

    if not rows:
        log("ERROR: No data rows found in Excel (only a header row exists).")
        sys.exit(1)

    return headers, rows


# ── WORD ──────────────────────────────────────────────────────────────────────

def _replace_in_paragraph(para, replacements):
    """
    Replace [key] occurrences in a paragraph.
    Merges all runs into the first run to handle cases where Word splits
    a placeholder across multiple runs (e.g. after spell-check or mid-word formatting).
    The first run's character formatting is preserved for the full paragraph text.
    """
    if not para.runs:
        return
    full = "".join(r.text or "" for r in para.runs)
    new = full
    for key, value in replacements.items():
        new = new.replace(f"[{key}]", value)
    if new != full:
        para.runs[0].text = new
        for r in para.runs[1:]:
            r.text = ""


def fill_template(template_path, row_data, output_path):
    """
    Open template_path, replace all [placeholder] tags with row_data values,
    save result to output_path.
    Handles body paragraphs, tables, and header/footer sections.
    """
    import shutil
    import tempfile
    from docx import Document

    # python-docx's zip layer treats [ ] as glob wildcards, so any brackets
    # in the original filename cause a "Package not found" error.
    # Copy to a temp file with a clean name before opening.
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    shutil.copy2(template_path, tmp_path)
    try:
        doc = Document(tmp_path)
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass

    # Body paragraphs
    for para in doc.paragraphs:
        _replace_in_paragraph(para, row_data)

    # Tables
    for table in doc.tables:
        for trow in table.rows:
            for cell in trow.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, row_data)

    # Headers and footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_paragraph(para, row_data)
        for para in section.footer.paragraphs:
            _replace_in_paragraph(para, row_data)

    doc.save(output_path)


# ── PDF CONVERSION ────────────────────────────────────────────────────────────

def convert_to_pdf(docx_path, pdf_path, log=print):
    """Convert .docx to .pdf via Microsoft Word (win32com). Word must be installed."""
    try:
        import win32com.client
        import pythoncom
    except ImportError:
        log("ERROR: pywin32 is not installed. Run: pip install pywin32")
        raise

    word = None
    pythoncom.CoInitialize()
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(os.path.abspath(docx_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # 17 = wdFormatPDF
        doc.Close(False)
    except Exception as e:
        log(f"  ERROR during PDF conversion: {e}")
        raise
    finally:
        if word:
            try:
                word.Quit()
            except Exception:
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


# ── MAIN PROCESS ──────────────────────────────────────────────────────────────

def process(template_path, to_pdf, progress_callback=None, log_callback=None):
    """
    Main entry point called by the launcher (or directly from __main__).
    template_path : absolute path to the .docx template
    to_pdf        : bool, whether to also produce a .pdf for each output
    progress_callback(current, total) : optional, for the launcher progress bar
    log_callback(msg)                 : optional, for the launcher log panel
    """
    log = log_callback if log_callback else print

    check_dependencies()

    if not os.path.isfile(template_path):
        log(f"ERROR: Template not found: {template_path}")
        sys.exit(1)

    template_name = os.path.splitext(os.path.basename(template_path))[0]

    log(f"Template   : {template_path}")
    log(f"Excel      : {EXCEL_PATH}")
    log(f"Output     : {OUTPUT_FOLDER}")
    log(f"PDF output : {'yes' if to_pdf else 'no'}")
    log("")

    headers, rows = read_excel(log)
    log(f"Placeholders : {', '.join(headers)}")
    log(f"Rows loaded  : {len(rows)}")
    log("")

    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    total = len(rows)

    for i, row in enumerate(rows, 1):
        file_label = (
            row.get("File name", "").strip()
        )
        if not file_label:
            log(f"  [{i}/{total}] SKIPPED — empty 'File name' cell")
            if progress_callback:
                progress_callback(i, total)
            continue
        file_label = (
            file_label
            .replace("/", "-")
            .replace("\\", "-")
            .replace(":", "-")
        )
        out_name = f"{template_name} - {file_label}"
        out_docx = os.path.join(OUTPUT_FOLDER, f"{out_name}.docx")

        fill_template(template_path, row, out_docx)
        log(f"  [{i}/{total}] {file_label} — .docx written")

        if to_pdf:
            out_pdf = os.path.join(OUTPUT_FOLDER, f"{out_name}.pdf")
            convert_to_pdf(out_docx, out_pdf, log)
            log(f"  [{i}/{total}] {file_label} — .pdf written")
            try:
                os.remove(out_docx)
            except Exception:
                pass

        if progress_callback:
            progress_callback(i, total)

    log(f"\nDone. {total} file(s) written to: {OUTPUT_FOLDER}")


# ── CLI ───────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python fill_template.py <template_path> [--pdf]")
        sys.exit(1)
    _template = sys.argv[1]
    _to_pdf = "--pdf" in sys.argv
    process(_template, _to_pdf)
